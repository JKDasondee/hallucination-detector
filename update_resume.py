from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT
import copy
import os
from lxml import etree

src = "C:/Users/Tesla/Desktop/Jay_Dasondee_Resume_2.docx"
out1 = "C:/Users/Tesla/Desktop/Jay_Dasondee_Resume_2.docx"
out2 = "C:/Users/Tesla/Desktop/Resumes/2026-April-Targeted/Jay_Dasondee_Resume_ML.docx"

doc = Document(src)

# 1. Update tagline (P3)
p3 = doc.paragraphs[3]
for r in p3.runs:
    r.text = ""
p3.runs[0].text = "ML/NLP Engineer | Operations Engineer | Support Systems & Automation | Full-Stack (Next.js/React/Python) | 44K+ Community Scaled | 100K+ Users Reached"

# 2. Add Hallucination Detector as first project (before Ambassador Hub = P16)
# Need to insert: title paragraph + 3 bullet paragraphs before P16

# Reference: P16 is project title, P17 is bullet
# We insert before P16's element in the XML tree

ref_title = doc.paragraphs[16]  # Ambassador Hub title
ref_bullet = doc.paragraphs[17]  # Ambassador Hub bullet

def clone_paragraph(p):
    new_el = copy.deepcopy(p._element)
    return new_el

# Create title paragraph by cloning P16's structure
title_el = clone_paragraph(ref_title)
# Clear all runs from clone
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for r in title_el.findall('.//w:r', nsmap):
    title_el.remove(r)

def make_run(text, bold=None, italic=None, size=None, color=None):
    r_el = etree.SubElement(title_el, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
    rpr = etree.SubElement(r_el, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if bold:
        etree.SubElement(rpr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
    if italic:
        etree.SubElement(rpr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i')
    if size:
        sz = etree.SubElement(rpr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz')
        sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(size))
        szcs = etree.SubElement(rpr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}szCs')
        szcs.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(size))
    if color:
        c = etree.SubElement(rpr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
        c.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', color)
    t = etree.SubElement(r_el, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    t.text = text
    if text.startswith(' ') or text.endswith(' ') or text.startswith('\t'):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r_el

# Title runs matching P16 format exactly
make_run("Hallucination Detector", bold=True, size=20, color=None)  # size 20 = 10pt in half-points... wait
# P16 R0 size=127000 EMU = 10pt. In python-docx XML, sz val is in half-points. 10pt = 20 half-points
make_run("  |", size=20, color="555555")
make_run("  github.com/JKDasondee/hallucination-detector", size=20, color="555555")
make_run("\t2026", italic=True, color="666666")

# Create 3 bullet paragraphs by cloning P17
bullets = [
    "Built end-to-end LLM hallucination detection pipeline: claim extraction (spaCy NER), evidence retrieval (ChromaDB + sentence-transformers), NLI entailment scoring (HuggingFace cross-encoder), and confidence aggregation (sklearn ensemble)",
    "FastAPI backend with Claude API integration, React frontend with claim-level color highlighting (green=verified, red=hallucinated, yellow=uncertain), deployed on Railway + Vercel",
    "Implemented RAG pipeline indexing 500K+ Wikipedia passages into ChromaDB vector store with all-MiniLM-L6-v2 embeddings for real-time evidence retrieval against generated claims",
]

bullet_els = []
for txt in bullets:
    bel = clone_paragraph(ref_bullet)
    for r in bel.findall('.//w:r', nsmap):
        bel.remove(r)
    r_el = etree.SubElement(bel, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
    t = etree.SubElement(r_el, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    t.text = txt
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    bullet_els.append(bel)

# Insert before P16
parent = ref_title._element.getparent()
idx = list(parent).index(ref_title._element)
for i, el in enumerate([title_el] + bullet_els):
    parent.insert(idx + i, el)

# 3. Update TECHNICAL SKILLS
# Re-read paragraphs after insertion (indices shifted by 4)
# Data & ML was P26, now P30
# Frameworks was P24, now P28

for p in doc.paragraphs:
    txt = p.text.strip()
    if txt.startswith("Data & ML:"):
        # Add spaCy, sentence-transformers, ChromaDB, HuggingFace, FAISS
        for r in p.runs:
            if not r.bold and r.text and "data pipeline design" in r.text:
                current = r.text
                additions = []
                for item in ["spaCy", "sentence-transformers", "ChromaDB", "HuggingFace", "FAISS"]:
                    if item not in txt:
                        additions.append(item)
                if additions:
                    r.text = current.replace("data pipeline design", ", ".join(additions) + ", data pipeline design")
                break
    elif txt.startswith("Frameworks:"):
        for r in p.runs:
            if not r.bold and "React" in r.text:
                if "FastAPI" not in txt:
                    r.text = r.text.replace("React,", "React, FastAPI,")
                break

# Save
doc.save(out1)
os.makedirs(os.path.dirname(out2), exist_ok=True)
doc.save(out2)
print("Done. Saved to both locations.")
